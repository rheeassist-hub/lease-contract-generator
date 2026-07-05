"""
임대차계약서 자동 생성기 - FastAPI 백엔드
한국 표준 임대차계약서 기반
"""

from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import json
import uuid
import os
from typing import Optional
import logging

# Logging 설정
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="Lease Contract Generator",
    description="한국 표준 임대차계약서 자동 생성 API",
    version="1.0.0"
)

# CORS 설정
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 배포 시 수정: ["https://your-domain.vercel.app"]
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==================== Models ====================

class Landlord(BaseModel):
    name: str
    phone: str
    address: str

class Tenant(BaseModel):
    name: str
    phone: str
    occupation: str

class ContractRequest(BaseModel):
    landlord: Landlord
    tenant: Tenant
    contract_type: str  # "jeonse" or "monthly"
    deposit: int  # 보증금 (KRW)
    monthly_rent: int  # 월세 (0 if jeonse)
    start_date: str  # YYYY-MM-DD
    end_date: str  # YYYY-MM-DD
    special_clauses: Optional[str] = None
    format: str = "docx"  # "docx", "pdf", "hwp"

# ==================== Utility Functions ====================

def format_number(num: int) -> str:
    """숫자를 한글 포맷으로 변환: 100000 -> '100,000'"""
    return f"{num:,}"

def korean_date(date_str: str) -> str:
    """2026-08-01 -> '2026년 8월 1일'"""
    try:
        d = datetime.strptime(date_str, "%Y-%m-%d")
        return f"{d.year}년 {d.month}월 {d.day}일"
    except:
        return date_str

def create_lease_docx(contract: ContractRequest) -> str:
    """DOCX 임대차계약서 생성"""
    doc = Document()
    
    # 제목
    title = doc.add_paragraph()
    title_run = title.add_run("임 대 차 계 약 서")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("(대법원 표준 임대차계약서)")
    subtitle_run.font.size = Pt(10)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 빈줄
    
    # 계약 당사자
    doc.add_heading("1. 계약 당사자", level=2)
    
    parties = doc.add_paragraph()
    parties.add_run(f"임대인(집주인): {contract.landlord.name} (전화: {contract.landlord.phone})\n")
    parties.add_run(f"임차인(세입자): {contract.tenant.name} (전화: {contract.tenant.phone}, 직업: {contract.tenant.occupation})\n")
    parties.add_run(f"임대주택소재지: {contract.landlord.address}")
    
    doc.add_paragraph()
    
    # 계약 목적물
    doc.add_heading("2. 임차목적물", level=2)
    
    property_info = doc.add_paragraph()
    property_info.add_run(f"주소: {contract.landlord.address}")
    
    doc.add_paragraph()
    
    # 계약금액
    doc.add_heading("3. 계약금액 및 납부방법", level=2)
    
    if contract.contract_type == "jeonse":
        amount = doc.add_paragraph()
        amount.add_run(f"보증금: ₩ {format_number(contract.deposit)}")
        doc.add_paragraph("계약금은 계약 체결일 현금으로 직접 지급하고, "
                         "명도일에 최종 잔금을 지급한다.")
    else:  # monthly
        amount = doc.add_paragraph()
        amount.add_run(f"보증금: ₩ {format_number(contract.deposit)}\n")
        amount.add_run(f"월세: ₩ {format_number(contract.monthly_rent)} (매월 1일 납부)")
        doc.add_paragraph("계약금은 계약 체결일에 현금으로 직접 지급한다.")
    
    doc.add_paragraph()
    
    # 계약기간
    doc.add_heading("4. 임차기간", level=2)
    
    period = doc.add_paragraph()
    start_korean = korean_date(contract.start_date)
    end_korean = korean_date(contract.end_date)
    period.add_run(f"계약기간: {start_korean} ～ {end_korean}")
    
    doc.add_paragraph()
    
    # 임대인의 의무
    doc.add_heading("5. 임대인의 의무", level=2)
    
    doc.add_paragraph("임대인은 다음과 같은 의무를 진다:")
    doc.add_paragraph("① 임차인에게 목적물을 약정한 기간 동안 사용·수익하게 할 의무")
    doc.add_paragraph("② 목적물의 사용·수익을 방해하지 아니할 의무")
    doc.add_paragraph("③ 목적물에 숨겨진 하자가 있는 경우 수리할 의무")
    
    doc.add_paragraph()
    
    # 임차인의 의무
    doc.add_heading("6. 임차인의 의무", level=2)
    
    doc.add_paragraph("임차인은 다음과 같은 의무를 진다:")
    doc.add_paragraph("① 약정한 임차료를 기한 내에 지급할 의무")
    doc.add_paragraph("② 목적물을 현상대로 유지할 의무")
    doc.add_paragraph("③ 임차인이 야기한 손해에 대한 배상 의무")
    
    doc.add_paragraph()
    
    # 특약사항
    doc.add_heading("7. 특약사항", level=2)
    
    if contract.special_clauses:
        special = doc.add_paragraph()
        special.add_run(contract.special_clauses)
    else:
        doc.add_paragraph("(특약사항 없음)")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 서명란
    doc.add_heading("8. 서명", level=2)
    
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = 'Light Grid Accent 1'
    
    # 임대인 서명
    sig_table.cell(0, 0).text = "임대인"
    sig_table.cell(0, 1).text = contract.landlord.name
    sig_table.cell(1, 0).text = "서명 / 인장"
    sig_table.cell(1, 1).text = "(서명 또는 인장)"
    sig_table.cell(2, 0).text = "임차인"
    sig_table.cell(2, 1).text = contract.tenant.name
    sig_table.cell(3, 0).text = "서명 / 인장"
    sig_table.cell(3, 1).text = "(서명 또는 인장)"
    
    doc.add_paragraph()
    
    footer = doc.add_paragraph()
    today = datetime.now().strftime("%Y년 %m월 %d일")
    footer.add_run(today)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 파일 저장
    output_dir = Path("./contracts")
    output_dir.mkdir(exist_ok=True)
    
    filename = f"lease_contract_{uuid.uuid4().hex[:8]}.docx"
    filepath = output_dir / filename
    
    doc.save(filepath)
    logger.info(f"Created contract: {filepath}")
    return str(filepath)

# ==================== Routes ====================

@app.get("/health")
def health_check():
    """헬스체크"""
    return {"status": "ok", "service": "lease-contract-generator"}

@app.get("/")
def root():
    """루트 엔드포인트"""
    return {
        "name": "Lease Contract Generator API",
        "version": "1.0.0",
        "docs": "/docs",
        "endpoints": {
            "health": "/health",
            "generate": "/api/generate-contract",
            "download": "/api/download/{filename}"
        }
    }

@app.post("/api/generate-contract")
def generate_contract(contract: ContractRequest):
    """임대차계약서 생성 및 다운로드"""
    try:
        logger.info(f"Generating contract for {contract.tenant.name}")
        
        # 현재는 DOCX만 지원
        if contract.format not in ["docx", "pdf", "hwp"]:
            raise HTTPException(status_code=400, detail="지원하지 않는 형식입니다. (현재 DOCX만 지원)")
        
        # DOCX 생성
        filepath = create_lease_docx(contract)
        
        return {
            "status": "success",
            "download_url": f"/api/download/{Path(filepath).name}",
            "filename": f"임대차계약서_{contract.start_date}.docx",
            "message": "계약서가 성공적으로 생성되었습니다."
        }
    except Exception as e:
        logger.error(f"Error generating contract: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/download/{filename}")
def download_contract(filename: str):
    """계약서 다운로드"""
    try:
        # 보안: 파일명 검증
        if ".." in filename or "/" in filename:
            raise HTTPException(status_code=400, detail="Invalid filename")
        
        filepath = Path("./contracts") / filename
        
        if not filepath.exists():
            logger.warning(f"File not found: {filepath}")
            raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")
        
        logger.info(f"Downloading contract: {filepath}")
        return FileResponse(
            filepath, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error downloading contract: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        app, 
        host="0.0.0.0", 
        port=int(os.getenv("PORT", 8000)),
        log_level="info"
    )
